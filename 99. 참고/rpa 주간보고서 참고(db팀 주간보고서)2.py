from docx import Document
from docx.shared import Inches
import os, fnmatch
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Length
from docx.oxml.ns import qn, nsdecls
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
import datetime



import smtplib
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email import encoders
# import schedule
import time

# 날짜 설정
today = datetime.datetime.today()
st_date = today - datetime.timedelta(7)
st_date = st_date.strftime('%Y.%m.%d')
en_date = today - datetime.timedelta(1)
en_date = en_date.strftime('%Y.%m.%d')
todayd = today.strftime('%Y.%m.%d')
today = today.strftime('%Y%m%d')
path = '/home/seonglf/img/'+today+'/'

# make docx file
def create_report():

    def get_fontSize(s):
        run = document.add_paragraph().add_run(s)
        font = run.font
        font.size = Pt(9)
        return font

    def get_title(s):
        run = document.add_paragraph().add_run(s)
        font = run.font
        font.name = 'MS Gothic(제목 한글)'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(0x1f,0x49,0x7d)
        return font

    document = Document()

    # add littlefox logo
    # document.add_picture('/home/seonglf/img/logo.png', width=Inches(2))
    logo = document.add_picture('/home/seonglf/img/logo.png', width=Inches(2))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add header
    h = document.add_heading('', 0)
    h.add_run('영어 운영 주간 모니터링 보고서').bold = True
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # h1.style = 'Heading 1'
    get_title('모니터링 기간 : '+st_date+' ~ '+en_date)
    get_title('대상 DB : orcl')
    get_title('오라클 버전 : 10.2.0.4')
    get_title('대상 서버 : 10.1.123.31')

    # create table
    records = (
        ('core', '4'),
        ('memory', '126G'),
        ('disk', '1.9T')
    )
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'os'
    hdr_cells[1].text = 'CentOS 6.10'

    #header cell
    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size= Pt(12)
    font.color.rgb = RGBColor(0x1f,0x49,0x7d)
    paragraph = hdr_cells[1].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size= Pt(12) # font size = 30
    font.color.rgb = RGBColor(0x1f,0x49,0x7d)
    # hdr_cells[2].text = 'Desc'
    for qty, id in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size= Pt(12) # font size = 30
        font.color.rgb = RGBColor(0x1f,0x49,0x7d)
        paragraph = row_cells[1].paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size= Pt(12) # font size = 30
        font.color.rgb = RGBColor(0x1f,0x49,0x7d)

    for cell in table.columns[0].cells:
        cell.width = Inches(2)

    for cell in table.columns[1].cells:
        cell.width = Inches(2.5)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(1)
    paragraph_format.space_before = Pt(1)
    document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)
    paragraph_format.space_after = Pt(10)

    # add images
    # path = "E:/giho/sherpa/report/20211129/"
    # images = os.listdir(path)
    images = fnmatch.filter(os.listdir(path), '*.png')
    images.sort()

    # 이미지 아래 코멘트 추가
    for image in images:
        i = document.add_paragraph()
        i.add_run(image[:-4]).bold=True
        if image[:1] == '1':
            get_fontSize('   - 지난 한주간 Tablespace 사용량을 보여주는 지표 입니다.')
            get_fontSize('   - 약 90% 이상이 되면 더 할당을 하여 가용 공간을 확보 할 예정 입니다.')
            get_fontSize('   - 미미하게 증가하고 있는 추세 입니다.')
        elif image[:1] == '2':
            get_fontSize('   - DB 서버내 디스크 사용량에 대한 지표 입니다.')
            get_fontSize('   - 각 디스크 파티션별 주초, 주말을 비교한 값 입니다.')
            get_fontSize('   - 전체 디스크 사용량이 90%에 근접 하면 공간 확보를 위한 조치를 합니다.')
        elif image[:1] == '3':
            get_fontSize('   - 지난 한주간의 DB서버 성능에 대한 종합 지표 입니다.')
            get_fontSize('   - 가장 위에 있는 그래프는 CPU 사용량으로써 전반적으로 비슷한 패턴이 보이면 이상이 없는 상태라고 할 수 있습니다. 파란색은 DB user의 CPU사용량, 핑크색은 시스템에서 사용하는 CPU량을 이야기 하고, 초록색이 DB에서 발생한 wait에서 사용한 량을 의미 합니다.')
            get_fontSize('   - 두번째는 메모리 사용량으로써 특별하게 튀는 현상이 뚜렷하게 관찰되지만 않으면 안정적이라고 볼 수 있습니다.')
            get_fontSize('   - 세번째는 Active Sessions 지표 입니다. 어떤 순간 실제 SQL이 실행되고 있는 Session의 수를 나타내는데 전체적으로 15이상을 넘지 않으면서 두꺼운 느낌이 아니라면 안정적이라고 판단 할 수 있습니다.')
            get_fontSize('   - 마지막 네번째는 Wait Class 지표 입니다. SQL이 실행 될 때 실제 CPU가 일을 하는 시간외에 더 드는 시간에 대한 지표인데, 푸른색이 많을수록 안정적이라고 볼 수 있습니다.')
        elif image[:1] == '4':
            get_fontSize('   - 지난 한주간을 시간별로 구분하여 Active Session 지표를 시각화한 Heatmap Trend 입니다. 색상이 연할수록 안정적이라고 볼 수 있습니다.')
        elif image[:1] == '5':
            get_fontSize('   - Machine 별 session을 카운트한 지표 입니다.')
            get_fontSize('   - apis로 시작하는 Machine은 APP을 통한 session을 의미하며, lf-sqlr로 시작하는 Machine들은 sql릴레이 서버를 통해 접근한 session을 의미 합니다. 나머지로 cms 관리자와 직접 DB에 접근하여 실행한 session등이 있습니다.')
        elif image[:1] == '6':
            get_fontSize('   - DB 성능에서 Wait에 대한 지표를 나타낸 그래프 입니다.')
            get_fontSize('   - 대표적으로 CPU time와 db file sequential read, log file sync 등등 있습니다. 말 그래도 CPU time은 어떤 SQL을 수행 할 때 CPU를 얼마나 사용 했는지를 나타내며, db file sequential read는 보통 인덱스 스캔, ROWID에 의한 테이블 스캔, 컨트롤 파일(Control file), 파일 헤더(File header)를 읽을 때 발생 합니다. 비교적 정상적인 wait event로 분류 할 수 있습니다.')
            get_fontSize('   - 대체적으로 푸른색 계열이 나타나고, 핑크색인 Lock이 많이 발생 하지 않으면 좋은 그래프로 볼 수 있습니다.')
        elif image[:1] == '7':
            get_fontSize('   - SQL 실행횟수에 따른 top10 SQL입니다. 여기에 나타난다고 해서 문제가 있다는 의미는 아닙니다. 어떤 SQL이 많이 실행되었는지를 나타냅니다.')
            get_fontSize('   - 관련 개발자분들은 해당 SQL을 참고 해주시면 됩니다.')
            get_fontSize('   - 특별한 문제가 있으면 DB팀에 수정을 요청하세요.')
        elif image[:1] == '8':
            get_fontSize('   - SQL 응답시간에 따른 top10 SQL입니다. 여기에 나타난다고 해서 문제가 있다는 의미는 아닙니다. 어떤 SQL이 수행되는데 오래 걸리는지를 나타냅니다.')
            get_fontSize('   - 관련 개발자분들은 해당 SQL을 참고 해주시면 됩니다.')
            get_fontSize('   - 특별한 문제가 있으면 DB팀에 수정을 요청하세요.')
        elif image[:1] == '9':
            get_fontSize('   - SQL Physical read에 따른 top10 SQL 입니다. 여기에 나타난다고 해서 문제가 있다는 의미는 아닙니다. 어떤 SQL이 메모리가 아닌 Disk에서 데이터를 많이 읽어오는지를 나타냅니다.')
            get_fontSize('   - 관련 개발자분들은 해당 SQL을 참고 해주시면 됩니다.')
            get_fontSize('   - 특별한 문제가 있으면 DB팀에 수정을 요청하세요.')
        document.add_picture(path+image, width=Inches(6))
        document.add_paragraph()

    document.add_page_break()
    document.save('/home/seonglf/report/'+today+'.docx')
#    document.save(path + 'demo.docx')
    print('complete create docx ' + time.strftime('%H' + ':' + '%M' + ':' + '%S'))

# 메일 발송
def send_mail():

    sd = st_date.split('.')
    ed = en_date.split('.')
    sdt = sd[0]+'년 '+sd[1]+'월 '+sd[2]+'일'
    edt = ed[0]+'년 '+ed[1]+'월 '+ed[2]+'일'

    # 세션 생성
    s = smtplib.SMTP('smtp.gmail.com', 587)

    # TLS 보안 시작
    s.starttls()

    # 로그인 인증
    s.login('seonglf@gmail.com', 'vupscpwwegldbtlg')

    # 보낼 메시지 내용
    msg = MIMEMultipart()
    msg['Subject'] = '영어 운영 주간 모니터링 보고서 {0} ~ {1}'.format(sdt, edt)
    body = '''
    안녕하세요.

    DB팀에서 {0} ~ {1}까지 영어 DB 주간 모니터링 보고서를 전달 드립니다.

    이 보고서는 해당 기간동안 셀파오라클에서 수집한 영어 DB의 리소스 사용 추이와 각 세션 및 SQL 별 각종 이벤트를 요약하여 만들었습니다.

    매주 월요일 아침 대표님 이하 IT본부 각 팀장님, 개발팀원 전원에게 자동 발송 되고 있습니다.

    보고서 내 이미지 파일은 확대하여 자세히 보실 수 있습니다. 언급된 Top SQL은 담당 개발자 분들께서 참고 하시면 되겠습니다.
    ※특별한 이슈가 있으면 DB팀에서 따로 노티를 드립니다 :)

    DB는 상시 모니터링 되고 있으며, 이슈가 있으면 최대한 빠르게 대응하여 조치 하고 있습니다.

    마지막으로 각 지표에 대하여 궁금한 점이 있으시면 DB팀에 문의 해 주시기 바랍니다.

    감사합니다.
    '''.format(sdt, edt)
    msg.attach(MIMEText(body, 'plain'))
    filename = '/home/seonglf/report/'+today+'.docx'
    attachment = open(filename, 'rb')
    part = MIMEBase('application', 'octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', "attachment; filename= " + 'Monitoring Report(' + today+').docx')
    msg.attach(part)
    people = ['seonglf@littlefox.com','dba.stricky@littlefox.com','yjyoon8425@littlefox.com',
            'killerza@littlefox.com',
            'kurum0107@littlefox.com',
            'syalittlefox@littlefox.com',
            'myeong.wp@littlefox.com',
            'yongberm.na@littlefox.com',
            'shinjm.little@littlefox.com',
            'irene1009jane@littlefox.com',
            'hyeonbae5809@littlefox.com',
            'psc5956@littlefox.com',
            'ahndoae@littlefox.com',
            'chakyunga91@littlefox.com',
            'tequila@littlefox.com',
            'han.js.dev@littlefox.com',
            'f1n4l34f@littlefox.com',
            'only340@littlefox.com',
            'jyan@littlefox.com',
            'root.joon@littlefox.com',
            'blackone0919@littlefox.com',
            'ceokim@littlefox.com',
            'ceoyang@littlefox.com',
            'mitsusumih@littlefox.com',
            'sozini94@littlefox.com']
    for p in people:
        s.sendmail("seonglf@gmail.com", p, msg.as_string())

    # 세션 종료
    s.quit()
    print('complete send e-mail ' + time.strftime('%H' + ':' + '%M' + ':' + '%S'))

create_report()
send_mail()

